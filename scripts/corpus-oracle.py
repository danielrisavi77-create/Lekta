#!/usr/bin/env python
"""
NEOVISAN ORAKUL nad .docx: druga implementacija koja mjeri iste cinjenice kao Lekta, ali svojim
putem, pa se njihova neslaganja mogu vidjeti.

ZASTO POSTOJI. Lekta je deterministicna, ali determinizam nije provjera: sustav moze biti dosljedno
u krivu. Vise prolaza ISTIM alatom zato nista ne dokazuje (FER pilot: 7/7 doslovnih citata, a 4 od 5
tvrdnji oboreno). Za dokaz treba orakul koji NE dijeli implementaciju s onim koga provjerava.

GRANICA KOJU OVAJ ORAKUL NE PRELAZI, i to je namjerno:
  - MJERI dokument vlastitim citacem (python-docx), bez ijednog retka Lektina parsera;
  - PRAVILO cita iz istog profila kao Lekta (verified-profiles.json). To NIJE propust: manifest
    provjerava mjeri li Lekta dokument tocno, a je li samo PRAVILO tocno provjerava verifikacijski
    lanac (citat, izvor, modalitet). Dvije razlicite tvrdnje, dva razlicita garda.
  - Ne ocjenjuje rad i ne dira dokument.

Izlaz je expected.findings za manifest dokaza: popis (checkId, expectFail, because).

Pokreni:  python scripts/corpus-oracle.py <docx> <profileId>
          python scripts/corpus-oracle.py --corpus
"""
from __future__ import annotations

import json
import os
import sys
from collections import Counter

try:
    import docx
except ImportError:
    print("python-docx nije instaliran (pip install python-docx)", file=sys.stderr)
    raise SystemExit(2)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CORPUS = os.environ.get("LEKTA_CORPUS_ROOT", "C:/Users/PC/LektaCorpus/corpus")
EMU_PER_CM = 360000
TOL_CM = 0.05
A4_W, A4_H = 21.0, 29.7


def cm(emu):
    return None if emu is None else round(emu / EMU_PER_CM, 2)


def profile_rules(profile_id):
    path = os.path.join(ROOT, "data/profiles/verified-profiles.json")
    with open(path, encoding="utf-8") as fh:
        for p in json.load(fh):
            if p.get("id") == profile_id:
                return p.get("rules") or {}
    return {}


def _doc_default_size(d):
    """Velicina iz docDefaults (w:rPrDefault/w:rPr/w:sz), u tockama."""
    try:
        el = d.styles.element
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        rpr = el.find(ns + "docDefaults/" + ns + "rPrDefault/" + ns + "rPr")
        if rpr is None:
            return None
        sz = rpr.find(ns + "sz")
        return float(sz.get(ns + "val")) / 2 if sz is not None else None
    except Exception:
        return None


def _style_size(style):
    """Velicina iz stila, uz penjanje po base_style lancu."""
    seen = 0
    while style is not None and seen < 12:
        if style.font is not None and style.font.size is not None:
            return style.font.size.pt
        style = getattr(style, "base_style", None)
        seen += 1
    return None


def _style_spacing(style):
    seen = 0
    while style is not None and seen < 12:
        pf = getattr(style, "paragraph_format", None)
        if pf is not None and pf.line_spacing is not None:
            return round(float(pf.line_spacing), 2)
        style = getattr(style, "base_style", None)
        seen += 1
    return None


def effective_size(par, run, doc_default):
    """Izravno na runu, pa stil odlomka, pa docDefaults. Isti redoslijed koji Word primjenjuje."""
    if run is not None and run.font.size is not None:
        return run.font.size.pt
    s = _style_size(par.style)
    return s if s is not None else doc_default


def _doc_default_spacing(d):
    """Prored iz docDefaults. `w:line` je u dvadesetinama retka, pa `auto` znaci line/240."""
    try:
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        ppr = d.styles.element.find(ns + "docDefaults/" + ns + "pPrDefault/" + ns + "pPr")
        if ppr is None:
            return None
        sp = ppr.find(ns + "spacing")
        if sp is None:
            return None
        line = sp.get(ns + "line")
        rule = sp.get(ns + "lineRule")
        if line is None or rule not in (None, "auto"):
            return None
        return round(float(line) / 240.0, 2)
    except Exception:
        return None


def effective_spacing(par, doc_default=None):
    """Izravno -> stil odlomka -> docDefaults. Bez zadnjeg koraka je 81 od 89 odlomaka ostajalo
    NEIZMJERENO, pa je "dominantan prored" nastajao iz osam iznimaka i lazno prijavljivao odstupanje
    (izmjereno na corpus-0075: docDefaults nosi w:line=360, dakle 1,5)."""
    if par.paragraph_format.line_spacing is not None:
        return round(float(par.paragraph_format.line_spacing), 2)
    s = _style_spacing(par.style)
    return s if s is not None else doc_default


def measure(path):
    """Cinjenice dokumenta, izmjerene python-docxom. Nikakva prosudba, samo brojke."""
    d = docx.Document(path)
    sec = d.sections[0] if d.sections else None
    fonts, sizes, spacings = Counter(), Counter(), Counter()
    # VAZNO: mjeri se STVARNO PRIMIJENJENA vrijednost, ne samo izravno formatiranje. Prvi prolaz je
    # brojao samo `run.font.size`, koji je zadan na malom dijelu runova (izmjereno: 15 od 248 na
    # corpus-0003), pa je "dominantna velicina" nastajala iz uzorka od 6 posto i lazno prijavljivala
    # odstupanje. Ispravno je razrijesiti lanac: izravno -> stil odlomka (uz base_style) -> docDefaults.
    doc_default = _doc_default_size(d)
    doc_default_sp = _doc_default_spacing(d)
    for p in d.paragraphs:
        if not p.text.strip():
            continue
        sp = effective_spacing(p, doc_default_sp)
        if sp is not None:
            spacings[sp] += 1
        for r in p.runs:
            if r.font.name:
                fonts[r.font.name] += 1
            es = effective_size(p, r, doc_default)
            if es is not None:
                sizes[es] += 1
        if not p.runs:
            es = effective_size(p, None, doc_default)
            if es is not None:
                sizes[es] += 1
    body_xml = d.element.body.xml
    # Broj stranice zna stajati u footeru PRVE ili PARNE stranice, i u zaglavlju, ne samo u zadanom
    # footeru. Prvi prolaz je gledao samo `s.footer` i zato lazno prijavljivao izostanak (izmjereno na
    # cetiri rada: PAGE je bio u `first_page_footer`).
    parts = []
    for s_ in d.sections:
        for name in ("footer", "first_page_footer", "even_page_footer",
                     "header", "first_page_header", "even_page_header"):
            part = getattr(s_, name, None)
            if part is not None and part._element is not None:
                parts.append(part._element.xml)
    footer_xml = " ".join(parts)
    return {
        "margins": None if sec is None else {
            "top": cm(sec.top_margin), "right": cm(sec.right_margin),
            "bottom": cm(sec.bottom_margin), "left": cm(sec.left_margin),
        },
        "pageWidthCm": None if sec is None else cm(sec.page_width),
        "pageHeightCm": None if sec is None else cm(sec.page_height),
        "dominantFont": fonts.most_common(1)[0][0] if fonts else None,
        "dominantSize": sizes.most_common(1)[0][0] if sizes else None,
        "dominantSpacing": spacings.most_common(1)[0][0] if spacings else None,
        "hasTocField": "TOC" in body_xml,
        "hasPageNumberField": "PAGE" in footer_xml,
        "paragraphs": len(d.paragraphs),
    }


def findings(m, rules):
    """Sto BI Lekta trebala naci, izvedeno iz NASIH mjerenja i profilnih pravila."""
    out = []

    def add(check_id, fail, because):
        out.append({"checkId": check_id, "expectFail": bool(fail), "because": because})

    want = rules.get("margins")
    if want and m["margins"] and rules.get("checkMargins") is not False:
        bad = [k for k in ("top", "right", "bottom", "left")
               if m["margins"].get(k) is not None and want.get(k) is not None
               and abs(m["margins"][k] - float(want[k])) > TOL_CM]
        add("page.margins", bool(bad),
            "izmjereno {} cm, profil trazi {} cm{}".format(
                m["margins"], want, ("; odstupa: " + ", ".join(bad)) if bad else ""))

    if rules.get("checkFont") and rules.get("font") and m["dominantFont"]:
        ok = m["dominantFont"] in list(rules["font"])
        add("format.font.dominant", not ok,
            "dominantan font je {}, profil trazi {}".format(m["dominantFont"], list(rules["font"])))

    if rules.get("checkSize") and rules.get("size") and m["dominantSize"] is not None:
        ok = float(m["dominantSize"]) in [float(x) for x in rules["size"]]
        add("format.size.body", not ok,
            "dominantna velicina je {} pt, profil trazi {}".format(m["dominantSize"], rules["size"]))

    if rules.get("checkSpacing") and rules.get("spacing") and m["dominantSpacing"] is not None:
        ok = abs(float(m["dominantSpacing"]) - float(rules["spacing"])) < 0.01
        add("format.spacing.body", not ok,
            "dominantan prored je {}, profil trazi {}".format(m["dominantSpacing"], rules["spacing"]))

    if rules.get("requireA4") and m["pageWidthCm"] and m["pageHeightCm"]:
        ok = abs(m["pageWidthCm"] - A4_W) < 0.2 and abs(m["pageHeightCm"] - A4_H) < 0.2
        add("page.size.a4", not ok,
            "stranica je {}x{} cm".format(m["pageWidthCm"], m["pageHeightCm"]))

    if rules.get("requireToc"):
        add("toc.present", not m["hasTocField"],
            "polje sadrzaja " + ("pronadjeno" if m["hasTocField"] else "NIJE pronadjeno"))

    if rules.get("requirePageNumbers"):
        add("page.numbers.present", not m["hasPageNumberField"],
            "PAGE polje u podnozju " + ("pronadjeno" if m["hasPageNumberField"] else "NIJE pronadjeno"))

    return out


def for_document(path, profile_id):
    m = measure(path)
    return {"measured": m, "findings": findings(m, profile_rules(profile_id))}


def _tmp_doc(build):
    """Sagradi .docx u privremenoj datoteci i vrati putanju."""
    import tempfile
    d = docx.Document()
    build(d)
    path = os.path.join(tempfile.mkdtemp(prefix="lekta-oracle-"), "t.docx")
    d.save(path)
    return path


def selftest():
    """Podmetnuti slucajevi za tri kvara koje je prvi prolaz nad korpusom imao. Svaki ima i
    NEGATIVNU kontrolu, jer mjera koja uvijek prolazi ne dokazuje nista."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    failures = []

    def check(name, got, want):
        if got != want:
            failures.append("PROMASAJ [{}]: ocekivano {}, dobiveno {}".format(name, want, got))

    # 1. Velicina iz STILA, kad je run nema. Prvi prolaz je brojao samo izravnu velicinu.
    def build_style_size(d):
        d.styles["Normal"].font.size = docx.shared.Pt(12)
        d.add_paragraph("tekst bez izravne velicine")
    m = measure(_tmp_doc(build_style_size))
    check("velicina iz stila", m["dominantSize"], 12.0)

    def build_direct_size(d):
        d.styles["Normal"].font.size = docx.shared.Pt(12)
        p = d.add_paragraph()
        p.add_run("izravno 20").font.size = docx.shared.Pt(20)
    m = measure(_tmp_doc(build_direct_size))
    check("izravna velicina nadjacava stil", m["dominantSize"], 20.0)

    # 2. Prored iz docDefaults. Prvi prolaz ga uopce nije gledao.
    def build_default_spacing(d):
        ppr = d.styles.element.find(ns + "docDefaults/" + ns + "pPrDefault/" + ns + "pPr")
        if ppr is None:
            raise RuntimeError("nema pPrDefault, test se ne moze postaviti")
        sp = ppr.find(ns + "spacing")
        if sp is None:
            from docx.oxml.ns import qn
            sp = ppr.makeelement(qn("w:spacing"), {})
            ppr.append(sp)
        sp.set(ns + "line", "360")
        sp.set(ns + "lineRule", "auto")
        d.add_paragraph("odlomak bez izravnog proreda")
    m = measure(_tmp_doc(build_default_spacing))
    check("prored iz docDefaults", m["dominantSpacing"], 1.5)

    # 3. PAGE polje u podnozju PRVE stranice, ne u zadanom.
    def build_first_page_footer(d):
        d.sections[0].different_first_page_header_footer = True
        fp = d.sections[0].first_page_footer
        fp.paragraphs[0].text = "PAGE"
        d.add_paragraph("tijelo")
    m = measure(_tmp_doc(build_first_page_footer))
    check("PAGE u podnozju prve stranice", m["hasPageNumberField"], True)

    def build_no_page(d):
        d.add_paragraph("tijelo bez ikakvog polja")
    m = measure(_tmp_doc(build_no_page))
    check("bez PAGE polja se NE prijavljuje", m["hasPageNumberField"], False)

    for f in failures:
        print("  " + f)
    print("negativne kontrole orakula: 5 slucajeva, promasaja: {}".format(len(failures)))
    return len(failures)


def main():
    if "--selftest" in sys.argv:
        raise SystemExit(1 if selftest() else 0)
    if "--corpus" in sys.argv:
        out_path = None
        if "--json" in sys.argv:
            j = sys.argv.index("--json")
            out_path = sys.argv[j + 1] if j + 1 < len(sys.argv) else None
        collected = {}
        total, unreadable = 0, 0
        for f in sorted(os.listdir(CORPUS)):
            if not f.endswith(".docx"):
                continue
            try:
                with open(os.path.join(CORPUS, f[:-5] + ".json"), encoding="utf-8") as fh:
                    side = json.load(fh)
            except Exception:
                continue
            pid = side.get("profileId")
            if not pid:
                continue
            try:
                res = for_document(os.path.join(CORPUS, f), pid)
            except Exception as exc:
                # Dokument koji se ne moze otvoriti NIJE tiho preskocen: sutnja bi izgledala kao uspjeh.
                unreadable += 1
                print("{:26} {:34} NECITLJIV ({})".format(f[:-5], pid, type(exc).__name__))
                continue
            total += 1
            collected[f[:-5]] = {"profileId": pid, **res}
            fails = sum(1 for x in res["findings"] if x["expectFail"])
            print("{:26} {:34} provjera {:2}  ocekivano padova {}".format(
                f[:-5], pid, len(res["findings"]), fails))
        print("\ndokumenata izmjereno: {} | necitljivih: {}".format(total, unreadable))
        if out_path:
            with open(out_path, "w", encoding="utf-8") as fh:
                json.dump(collected, fh, ensure_ascii=False, indent=1)
            print("zapisano: " + out_path)
        return
    if len(sys.argv) < 3:
        print(__doc__)
        raise SystemExit(2)
    print(json.dumps(for_document(sys.argv[1], sys.argv[2]), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
