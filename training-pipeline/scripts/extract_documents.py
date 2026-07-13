#!/usr/bin/env python3
"""Extract text and optionally create DOCX working copies from PDFs."""
from __future__ import annotations

import argparse
from pathlib import Path

from training_pipeline_bootstrap import add_pipeline_to_path

add_pipeline_to_path()
from lib import read_jsonl, write_jsonl  # noqa: E402


def pdf_text(path: Path) -> str:
    import fitz

    with fitz.open(path) as document:
        return "\n".join(page.get_text() for page in document)


def docx_text(path: Path) -> str:
    from docx import Document

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def pdf_to_docx(path: Path, output_dir: Path, document_id: str) -> Path | None:
    try:
        from pdf2docx import Converter
    except ImportError:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    target = output_dir / f"{document_id}.docx"
    converter = Converter(str(path))
    try:
        converter.convert(str(target))
    finally:
        converter.close()
    return target if target.exists() else None


def records(input_path: Path, docx_dir: Path, convert_pdf: bool):
    for record in read_jsonl(input_path):
        path = Path(record["_localPath"])
        try:
            extracted = pdf_text(path) if path.suffix.lower() == ".pdf" else docx_text(path)
            record["extractedText"] = extracted
            record["extractionStatus"] = "ok" if extracted.strip() else "ocr_required"
            if path.suffix.lower() == ".docx":
                record["_analysisDocxPath"] = str(path)
            elif convert_pdf and extracted.strip():
                converted = pdf_to_docx(path, docx_dir, record["id"])
                if converted:
                    record["_analysisDocxPath"] = str(converted)
        except Exception as exc:
            record["extractionStatus"] = "error"
            record["extractionError"] = type(exc).__name__
        yield record


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--docx-dir", required=True, type=Path)
    parser.add_argument("--convert-pdf", action="store_true")
    args = parser.parse_args()
    count = write_jsonl(args.output, records(args.manifest, args.docx_dir, args.convert_pdf))
    print(f"Extracted {count} documents")


if __name__ == "__main__":
    main()
